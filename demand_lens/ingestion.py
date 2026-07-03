"""Validated, structure-aware extraction for user-uploaded knowledge files."""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path


MAX_FILE_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx"}
EXPECTED_MIME_TYPES = {
    ".txt": {"text/plain"},
    ".md": {"text/markdown", "text/plain", "text/x-markdown"},
    ".markdown": {"text/markdown", "text/plain", "text/x-markdown"},
    ".html": {"text/html"},
    ".htm": {"text/html"},
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}


@dataclass(frozen=True)
class ExtractedSection:
    content: str
    section_path: str
    page_number: int | None = None


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    mime_type: str
    byte_size: int
    extraction_method: str
    sections: list[ExtractedSection]

    @property
    def text(self) -> str:
        return "\n\n".join(section.content for section in self.sections)


class UploadValidationError(ValueError):
    pass


class _ReadableHTML(HTMLParser):
    BLOCKS = {"p", "div", "section", "article", "li", "h1", "h2", "h3", "h4", "h5", "h6", "tr"}

    def __init__(self):
        super().__init__()
        self.parts: list[str] = []
        self.ignored = 0

    def handle_starttag(self, tag: str, attrs):
        if tag in {"script", "style", "noscript"}:
            self.ignored += 1
        elif tag in self.BLOCKS:
            if self.parts:
                self.parts.append("\n\n")
            if re.fullmatch(r"h[1-6]", tag):
                self.parts.append("#" * int(tag[1]) + " ")

    def handle_endtag(self, tag: str):
        if tag in {"script", "style", "noscript"} and self.ignored:
            self.ignored -= 1
        elif tag in self.BLOCKS:
            self.parts.append("\n\n")

    def handle_data(self, data: str):
        if not self.ignored:
            self.parts.append(data)


def _safe_filename(filename: str) -> str:
    name = Path(filename or "upload").name
    name = re.sub(r"[^A-Za-z0-9._ -]", "_", name).strip(". ")
    if not name:
        raise UploadValidationError("The uploaded file needs a valid filename")
    return name[:200]


def _decode_text(data: bytes) -> str:
    if b"\x00" in data[:4096]:
        raise UploadValidationError("Binary content is not valid text")
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise UploadValidationError("Text encoding is not supported")


def _markdown_sections(text: str, fallback: str) -> list[ExtractedSection]:
    sections: list[ExtractedSection] = []
    headings: list[tuple[int, str]] = []
    current: list[str] = []

    def flush() -> None:
        content = "\n".join(current).strip()
        if content:
            path = " > ".join(title for _level, title in headings) or fallback
            sections.append(ExtractedSection(content=content, section_path=path))

    for line in text.splitlines():
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            flush()
            current = []
            level = len(match.group(1))
            title = match.group(2).strip()
            headings = [(old_level, old_title) for old_level, old_title in headings if old_level < level]
            headings.append((level, title))
        else:
            current.append(line)
    flush()
    return sections or [ExtractedSection(text.strip(), fallback)]


def validate_and_extract(filename: str, data: bytes, declared_mime: str | None = None) -> ExtractedDocument:
    name = _safe_filename(filename)
    extension = Path(name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise UploadValidationError(f"Unsupported file type: {extension or 'unknown'}")
    normalized_mime = (declared_mime or "").split(";", 1)[0].strip().lower()
    if normalized_mime and normalized_mime != "application/octet-stream" and normalized_mime not in EXPECTED_MIME_TYPES[extension]:
        raise UploadValidationError(
            f"Declared MIME type {normalized_mime} does not match the {extension} extension"
        )
    if not data:
        raise UploadValidationError("The uploaded file is empty")
    if len(data) > MAX_FILE_BYTES:
        raise UploadValidationError("File exceeds the 10 MB upload limit")

    if extension == ".pdf":
        if not data.startswith(b"%PDF-"):
            raise UploadValidationError("The file extension is PDF but its signature is invalid")
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise UploadValidationError("PDF extraction dependency is unavailable") from exc
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise UploadValidationError("Password-protected PDFs are not supported")
        sections = []
        for page_number, page in enumerate(reader.pages, 1):
            content = (page.extract_text() or "").strip()
            if content:
                sections.append(ExtractedSection(content, f"Page {page_number}", page_number))
        mime = "application/pdf"
        method = "pypdf"
    elif extension == ".docx":
        if not zipfile.is_zipfile(io.BytesIO(data)) or b"[Content_Types].xml" not in data[:5000]:
            # The central directory may not be near the front, so verify its member explicitly.
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as archive:
                    if "[Content_Types].xml" not in archive.namelist():
                        raise UploadValidationError("The DOCX package signature is invalid")
            except zipfile.BadZipFile as exc:
                raise UploadValidationError("The DOCX package signature is invalid") from exc
        try:
            from docx import Document
        except ImportError as exc:
            raise UploadValidationError("DOCX extraction dependency is unavailable") from exc
        document = Document(io.BytesIO(data))
        lines = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name.lower() if paragraph.style else ""
            level_match = re.search(r"heading\s+(\d)", style)
            lines.append(f"{'#' * int(level_match.group(1))} {text}" if level_match else text)
        for table in document.tables:
            lines.append("\n" + "\n".join(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows))
        sections = _markdown_sections("\n\n".join(lines), name)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        method = "python-docx"
    elif extension in {".html", ".htm"}:
        parser = _ReadableHTML()
        parser.feed(_decode_text(data))
        text = re.sub(r"[ \t]+", " ", "".join(parser.parts))
        text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
        sections = _markdown_sections(text, name)
        mime = "text/html"
        method = "html-parser"
    else:
        text = _decode_text(data).strip()
        sections = _markdown_sections(text, name) if extension in {".md", ".markdown"} else [ExtractedSection(text, name)]
        mime = "text/markdown" if extension in {".md", ".markdown"} else "text/plain"
        method = "structured-text"

    if not sections or sum(len(section.content) for section in sections) < 20:
        raise UploadValidationError("No usable text could be extracted from the file")
    return ExtractedDocument(name, mime, len(data), method, sections)
